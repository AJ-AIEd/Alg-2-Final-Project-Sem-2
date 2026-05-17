from pathlib import Path

from docx import Document
from docx.enum.section import WD_ORIENT
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt, RGBColor
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from reportlab.lib import colors
from reportlab.lib.enums import TA_CENTER, TA_LEFT
from reportlab.lib.pagesizes import landscape, letter
from reportlab.lib.styles import ParagraphStyle, getSampleStyleSheet
from reportlab.lib.units import inch
from reportlab.platypus import (
    SimpleDocTemplate,
    Paragraph,
    Spacer,
    Table,
    TableStyle,
    PageBreak,
    Flowable,
    KeepTogether,
)

ROOT = Path(__file__).resolve().parents[1]
OUT = ROOT / "downloads"
OUT.mkdir(exist_ok=True)
PAGE_SIZE = landscape(letter)
PAGE_W, PAGE_H = PAGE_SIZE

INK = colors.HexColor("#111827")
CHARCOAL = colors.HexColor("#334155")
MUTED = colors.HexColor("#64748B")
LINE = colors.HexColor("#DBE2EA")
SOFT = colors.HexColor("#F8FAFC")
GREEN = colors.HexColor("#15803D")
GREEN_SOFT = colors.HexColor("#EDF8F1")
BLUE = colors.HexColor("#2563EB")
BLUE_SOFT = colors.HexColor("#EDF4FF")
PURPLE = colors.HexColor("#7E22CE")
PURPLE_SOFT = colors.HexColor("#F6EFFF")
WARM = colors.HexColor("#F5F2ED")
WHITE = colors.white
RED = colors.HexColor("#B45353")

GROUP_JOURNAL_MODEL = (
    "The physical journal captures the group's evolving mathematical thinking. "
    "It should look used: sketched on, annotated, revised, taped into, crossed out, and questioned."
)

INDIVIDUAL_EVIDENCE_MODEL = (
    "Your individual digital evidence captures your own mathematical reasoning inside the group investigation. "
    "This may include personal graph notes, recalculations, interpretations, AI critiques, limitation reflections, "
    "and explanations of how your thinking changed."
)

PHASES = [
    ("Class 1", "Entering the System", "What relationships are we beginning to notice?", MUTED, [
        ("System map canvas", "Sketch variables, links, pressures, unknowns, and possible feedback loops.", True),
        ("Variable table and measurement notes", "Variable | Units | Why it matters | What is hard to measure. Which variables seem politically, socially, or physically difficult to measure?", False),
        ("Sketch before technology", "Predict the shape of one relationship before using Desmos, Sheets, or another tool. What relationship seems obvious at first but may be misleading?", True),
        ("Vlog notes", "Discuss uncertainty, early assumptions, possible misleading patterns, and variables that may be hard to measure reliably.", False),
    ]),
    ("Class 2", "Modeling Change", "How is the system changing over time?", GREEN, [
        ("Discrete vs continuous reasoning", "Does your context change in steps, continuously, or both? Justify the representation that best fits your system.", False),
        ("Pattern behavior", "Does the system behave additively, multiplicatively, unpredictably, or in phases? What evidence supports this?", False),
        ("Recursive and explicit formulas", "Write useful formulas if the pattern supports them. Annotate what each term, parameter, and assumption means in your context.", True),
        ("Accumulated change", "What does accumulated change represent in your system? Would accumulation realistically continue forever? Why or why not?", False),
        ("Meaningful time scales", "Choose multiple time ranges that matter for your investigation. How does the model behave differently across those scales?", True),
        ("Pattern breakdown", "When does the pattern stop behaving consistently? What evidence would make you revise or reject it?", False),
    ]),
    ("Class 3", "Acceleration, Decay, and Thresholds", "When does repeated change become accelerated or compounding?", BLUE, [
        ("Exponential vs linear", "What evidence suggests the change depends on the current amount rather than a constant added amount?", False),
        ("Sampled values vs continuous model", "How does continuous exponential modeling change the interpretation compared to sampled or discrete terms in your context?", False),
        ("Build the model", "Define starting value, growth or decay factor, variables, assumptions, and realistic domain. Which assumption is doing the most work?", True),
        ("Logarithmic threshold", "Use logarithms when you need to solve for an unknown time or threshold. What does that threshold mean in the real system?", True),
        ("Rate or assumption comparison", "Compare multiple realistic rates or assumptions relevant to your investigation. Which assumptions cause the prediction to change dramatically?", True),
        ("Model breakdown", "Where does exponential behavior stop being realistic? What real-world forces might interrupt it? What hidden variables matter?", False),
    ]),
    ("Class 4", "Constraints, Capacity, and Cost", "How do physical and economic constraints reshape the system?", PURPLE, [
        ("Flat blueprint", "Original dimensions, square cutout variable, folding lines, units in centimeters.", True),
        ("Folded storage structure", "Height, interior dimensions, usable space, loading floor.", True),
        ("Volume polynomial", "Factored form, standard form, and why multiplying dimensions creates a cubic.", False),
        ("Graph annotation", "Intercepts, multiplicity, extrema, domain, unrealistic regions, physical meaning.", True),
        ("End behavior and structure", "Degree, even/odd, leading coefficient, as x → ∞ and as x → −∞.", False),
        ("Polynomial division", "How could division help analyze storage, packing, capacity, grouping, or remaining space? Interpret the quotient and remainder physically.", True),
        ("Constraint revision", "Change a realistic physical or economic constraint from your context. What tradeoff appears, and which graph regions become impossible?", True),
    ]),
    ("Class 5", "Revision, Synthesis, and Communication", "What story does our mathematics tell?", INK, [
        ("Synthesis table", "Lens | What did it reveal? | Where did it fail?", False),
        ("Where models agreed or conflicted", "Compare mathematical lenses and assumptions.", False),
        ("Final mathematical claim", "Use evidence, assumptions, limitations, tradeoffs, and implications. What remained uncertain even after modeling?", False),
        ("Communication storyboard", "Choose a format. Plan how mathematical evidence, revisions, and uncertainty will be visible.", True),
        ("If we repeated this investigation", "With better data, what would change? Which assumption, measurement, or model would you revisit first?", False),
        ("Individual artifact snapshot", "Revision, hidden-variable analysis, limitation critique, or r/R² interpretation. A high R² can still mislead; correlation can support a claim without proving causation.", False),
    ]),
]

styles = getSampleStyleSheet()
styles.add(ParagraphStyle("Kicker", parent=styles["Normal"], fontName="Helvetica-Bold", fontSize=8.5, leading=11, textColor=MUTED, spaceAfter=7, uppercase=True))
styles.add(ParagraphStyle("TitleBig", parent=styles["Title"], fontName="Helvetica-Bold", fontSize=39, leading=42, textColor=INK, alignment=TA_LEFT, spaceAfter=14))
styles.add(ParagraphStyle("SectionTitle", parent=styles["Heading1"], fontName="Helvetica-Bold", fontSize=25, leading=30, textColor=INK, spaceAfter=10))
styles.add(ParagraphStyle("Subhead", parent=styles["Heading2"], fontName="Helvetica-Bold", fontSize=14, leading=18, textColor=INK, spaceAfter=7))
styles.add(ParagraphStyle("Body", parent=styles["BodyText"], fontName="Helvetica", fontSize=10.5, leading=15.5, textColor=CHARCOAL, spaceAfter=7))
styles.add(ParagraphStyle("Small", parent=styles["BodyText"], fontName="Helvetica", fontSize=9.2, leading=13.2, textColor=CHARCOAL, spaceAfter=4))
styles.add(ParagraphStyle("Tiny", parent=styles["BodyText"], fontName="Helvetica", fontSize=8.1, leading=11.2, textColor=CHARCOAL, spaceAfter=3))
styles.add(ParagraphStyle("Quote", parent=styles["BodyText"], fontName="Helvetica-Bold", fontSize=16, leading=22, textColor=INK, leftIndent=12, borderColor=INK, borderWidth=0, borderPadding=0, spaceAfter=7))
styles.add(ParagraphStyle("CenterTitle", parent=styles["Title"], fontName="Helvetica-Bold", fontSize=20, leading=25, alignment=TA_CENTER, textColor=INK, spaceAfter=18))
styles.add(ParagraphStyle("RubricBody", parent=styles["BodyText"], fontName="Helvetica", fontSize=9.4, leading=12.8, textColor=CHARCOAL, spaceAfter=3))
styles.add(ParagraphStyle("RubricLevel", parent=styles["Heading3"], fontName="Helvetica-Bold", fontSize=11.2, leading=14, textColor=INK, spaceAfter=5))


def p(text, style="Body"):
    return Paragraph(str(text).replace("\n", "<br/>"), styles[style])


def spacer(h=12):
    return Spacer(1, h)


def doc(path, title):
    return SimpleDocTemplate(
        str(path),
        pagesize=PAGE_SIZE,
        rightMargin=0.58 * inch,
        leftMargin=0.58 * inch,
        topMargin=0.48 * inch,
        bottomMargin=0.52 * inch,
        title=title,
    )


def footer(canvas, document):
    canvas.saveState()
    canvas.setStrokeColor(colors.HexColor("#E5EAF0"))
    canvas.setLineWidth(0.6)
    canvas.line(document.leftMargin, 28, PAGE_W - document.rightMargin, 28)
    canvas.setFont("Helvetica", 8)
    canvas.setFillColor(MUTED)
    canvas.drawString(document.leftMargin, 15, "Systems Under Pressure | Algebra 2 Final Project")
    canvas.drawRightString(PAGE_W - document.rightMargin, 15, str(document.page))
    canvas.restoreState()


class SystemMap(Flowable):
    def __init__(self, width=250, height=210):
        super().__init__()
        self.width = width
        self.height = height

    def draw(self):
        c = self.canv
        cx, cy = self.width * 0.48, self.height * 0.5
        nodes = [
            ("Climate", GREEN, -82, 66), ("Transport", BLUE, 56, 82), ("Storage", PURPLE, 111, 10),
            ("Exports", PURPLE, 94, -72), ("Food loss", RED, -35, -82), ("Demand", BLUE, -96, -22), ("Supply", GREEN, -105, 26),
        ]
        c.setStrokeColor(colors.HexColor("#BAC5D2"))
        c.setLineWidth(1.1)
        for _, _, dx, dy in nodes:
            c.line(cx, cy, cx + dx, cy + dy)
        c.setFillColor(WHITE)
        c.setStrokeColor(colors.HexColor("#CBD5E1"))
        c.circle(cx, cy, 34, stroke=1, fill=1)
        c.setFillColor(INK)
        c.setFont("Helvetica-Bold", 10)
        c.drawCentredString(cx, cy + 3, "Food")
        c.drawCentredString(cx, cy - 10, "System")
        for name, col, dx, dy in nodes:
            x, y = cx + dx, cy + dy
            c.setFillColor(WHITE)
            c.setStrokeColor(col)
            c.setLineWidth(1.4)
            c.circle(x, y, 24, stroke=1, fill=1)
            c.setFillColor(col)
            c.setFont("Helvetica-Bold", 7.4)
            c.drawString(x + 28, y - 3, name)
        c.setFillColor(MUTED)
        c.setFont("Helvetica", 8)
        c.drawCentredString(cx, cy - 55, "interconnected pressures")


class Workspace(Flowable):
    def __init__(self, label="student thinking space", height=220, grid=False):
        super().__init__()
        self.label = label
        self.height = height
        self.grid = grid
        self.width = 1

    def wrap(self, availWidth, availHeight):
        self.width = availWidth
        return availWidth, self.height

    def draw(self):
        c = self.canv
        c.setStrokeColor(colors.HexColor("#AEBBCC"))
        c.setFillColor(WHITE)
        c.roundRect(0, 0, self.width, self.height, 8, stroke=1, fill=1)
        c.setFont("Helvetica-Bold", 8.5)
        c.setFillColor(MUTED)
        c.drawString(14, self.height - 22, self.label)
        c.setStrokeColor(colors.HexColor("#E7ECF3"))
        c.setLineWidth(0.45)
        if self.grid:
            step = 28
            x = step
            while x < self.width:
                c.line(x, 14, x, self.height - 36)
                x += step
            y = 24
            while y < self.height - 36:
                c.line(14, y, self.width - 14, y)
                y += step
        else:
            y = self.height - 48
            while y > 16:
                c.line(16, y, self.width - 16, y)
                y -= 22


def card(title, body, accent=INK, fill=WHITE, body_style="Body"):
    content = [[p(title, "Subhead")], [p(body, body_style)]]
    t = Table(content, colWidths=[None], hAlign="LEFT")
    t.setStyle(TableStyle([
        ("BACKGROUND", (0, 0), (-1, -1), fill),
        ("BOX", (0, 0), (-1, -1), 0.8, LINE),
        ("LINEABOVE", (0, 0), (-1, 0), 5, accent),
        ("LEFTPADDING", (0, 0), (-1, -1), 15),
        ("RIGHTPADDING", (0, 0), (-1, -1), 15),
        ("TOPPADDING", (0, 0), (-1, -1), 13),
        ("BOTTOMPADDING", (0, 0), (-1, -1), 13),
        ("VALIGN", (0, 0), (-1, -1), "TOP"),
    ]))
    return t


def card_grid(items, cols=3):
    rows, row = [], []
    for title, body, accent, fill in items:
        row.append(card(title, body, accent, fill))
        if len(row) == cols:
            rows.append(row)
            row = []
    if row:
        while len(row) < cols:
            row.append(Spacer(1, 1))
        rows.append(row)
    usable = PAGE_W - 0.58 * inch * 2
    col_width = usable / cols
    t = Table(rows, colWidths=[col_width] * cols, hAlign="LEFT")
    t.setStyle(TableStyle([
        ("VALIGN", (0, 0), (-1, -1), "TOP"),
        ("LEFTPADDING", (0, 0), (-1, -1), 6),
        ("RIGHTPADDING", (0, 0), (-1, -1), 6),
        ("TOPPADDING", (0, 0), (-1, -1), 7),
        ("BOTTOMPADDING", (0, 0), (-1, -1), 7),
    ]))
    return t


def section(kicker, title, body=None):
    out = [p(kicker.upper(), "Kicker"), p(title, "SectionTitle")]
    if body:
        out.append(p(body, "Body"))
    out.append(spacer(10))
    return out


def title_story(subtitle):
    return [
        Table([[[
            p("ALGEBRA 2 CUMULATIVE SEMESTER FINAL PROJECT", "Kicker"),
            p("Systems<br/>Under Pressure", "TitleBig"),
            p("The Mathematics of Food Systems", "SectionTitle"),
            spacer(18),
            card('"All models are wrong, but some are useful."', "George Box", INK, WHITE),
            spacer(28),
            p(subtitle, "Body"),
        ], SystemMap(260, 220)]], colWidths=[PAGE_W * 0.52, PAGE_W * 0.30], hAlign="LEFT", style=TableStyle([
            ("VALIGN", (0, 0), (-1, -1), "MIDDLE"),
            ("LEFTPADDING", (0, 0), (-1, -1), 0),
            ("RIGHTPADDING", (0, 0), (-1, -1), 10),
        ])),
        PageBreak(),
    ]


RUBRIC_STANDARDS = [
    ("1", "Sequences and Series", "Content Standard", GREEN, GREEN_SOFT,
     [
         "Selects and justifies appropriate sequence and series models using strong mathematical reasoning. Clearly distinguishes arithmetic vs geometric, recursive vs explicit, finite vs infinite, and discrete vs continuous representations when relevant. Explains what terms, parameters, and cumulative sums mean in context. Predictions, recalculations, and interpretations are thoughtful, realistic, and well-supported.",
         "Uses appropriate sequence and series models with generally accurate reasoning. Explains arithmetic vs geometric patterns and recursive vs explicit forms with minor errors or missing detail. Interpretations and predictions are mostly reasonable and connected to context.",
         "Shows partial understanding of sequences and series. Some model choices or explanations are inaccurate, incomplete, or weakly justified. Connections between formulas, graphs, and context may be unclear or inconsistent.",
         "Shows limited understanding of sequences and series. Models are inappropriate, incomplete, or unsupported. Explanations and interpretations are mostly missing or incorrect.",
     ]),
    ("2", "Exponentials and Logarithms", "Content Standard", BLUE, BLUE_SOFT,
     [
         "Selects and justifies exponential models appropriately. Clearly explains growth and decay, thresholds, growth factors, and logarithmic reasoning in context. Thoughtfully compares exponential behavior to linear or geometric behavior when appropriate. Solutions are interpreted realistically with strong contextual understanding and critique of limitations.",
         "Uses exponential and logarithmic reasoning appropriately with mostly accurate explanations. Shows general understanding of growth factors, thresholds, and logarithmic solving. Interpretations are usually reasonable with only minor errors or missing detail.",
         "Shows partial understanding of exponentials and logarithms. Makes multiple errors in solving, interpretation, or model choice. Explanations may be vague, inconsistent, or disconnected from context.",
         "Shows limited understanding of exponential or logarithmic reasoning. Work is incomplete, incorrect, or lacks meaningful interpretation.",
     ]),
    ("3", "Polynomials", "Content Standard", PURPLE, PURPLE_SOFT,
     [
         "Creates and analyzes polynomial models with strong mathematical reasoning. Clearly explains degree, leading coefficient, end behavior, intercepts, multiplicity, extrema, and realistic domain restrictions in context. Thoughtfully compares factored and standard form and explains what each representation reveals. Uses notation such as x → ∞ and x → −∞ correctly and meaningfully.",
         "Analyzes polynomial models accurately with generally clear explanations. Shows understanding of graph behavior, structure, and polynomial forms with only minor errors or missing detail. Most interpretations are connected to context.",
         "Shows partial understanding of polynomial structure and graph behavior. Multiple errors or incomplete explanations weaken interpretation and analysis. Contextual meaning may be unclear or inconsistent.",
         "Shows limited understanding of polynomial reasoning or graph behavior. Analysis is incomplete, incorrect, or unsupported by meaningful explanation.",
     ]),
    ("4", "Mathematical Communication and Interpretation", "Skill Grade", INK, SOFT,
     [
         "Mathematical thinking is communicated clearly and thoughtfully through explanations, graphs, annotations, tables, discussion, and contextual interpretation. Strong connections among graph, algebra, assumptions, and context are consistently visible. Reasoning is organized, precise, and easy to follow.",
         "Communication is generally clear and understandable. Most mathematical ideas are connected to graphs, algebra, or context appropriately, though some explanations may lack precision or depth.",
         "Explanations are inconsistent, incomplete, or difficult to follow. Connections between mathematics and context are weak, unclear, or only partially developed.",
         "Mathematical thinking is minimally communicated. Explanations, interpretations, and connections are mostly missing or unclear.",
     ]),
    ("5", "Comparing Models, Assumptions, and Limitations", "Skill Grade", MUTED, SOFT,
     [
         "Thoughtfully compares multiple models, assumptions, and representations such as discrete vs continuous, recursive vs explicit, linear vs exponential, or different polynomial structures. Critiques realism, hidden variables, tradeoffs, assumptions, and unreasonable predictions with strong insight. Revision strengthens the mathematical investigation in meaningful ways.",
         "Compares models and assumptions appropriately with some explanation of limitations or realism. Shows general understanding of how assumptions affect predictions and interpretations. Some revision or reconsideration of ideas is visible.",
         "Comparisons are limited, superficial, or weakly justified. Explanations of assumptions, realism, or limitations are incomplete or inconsistent. Revision is minimal or mostly procedural.",
         "Shows little or no meaningful comparison of models, assumptions, or limitations. Revision and critique are mostly absent.",
     ]),
    ("6", "Algebraic Precision and Flexible Number Sense", "Skill Grade", INK, SOFT,
     [
         "Solves equations accurately using appropriate algebraic methods such as factoring, logarithms, inverse operations, substitution, graph interpretation, and algebraic manipulation. Fractions, decimals, percentages, negative numbers, and units are used fluently and precisely. Work is organized, justified, and consistently checked for reasonableness.",
         "Uses appropriate algebraic methods with mostly accurate calculations and organization. Rational numbers and units are generally used correctly, with only minor errors or omissions. Most answers are reasonable and interpreted appropriately.",
         "Uses some correct methods but makes multiple algebraic, numerical, or organizational errors. Precision, units, or interpretation may be inconsistent or incomplete.",
         "Shows little or no valid algebraic progress. Calculations, precision, organization, or interpretation are mostly incorrect or missing.",
     ]),
]


def journal_page(story, label, prompt, workspace_label="open thinking space", grid=False):
    story += section(label, prompt)
    story.append(Table([[Workspace(workspace_label, 285, grid), card("Reflection prompt", "What changed in your thinking? What assumption is doing the most work? What evidence would make this stronger?", INK, SOFT)]], colWidths=[PAGE_W - 260, 145], style=TableStyle([
        ("VALIGN", (0, 0), (-1, -1), "TOP"),
        ("LEFTPADDING", (0, 0), (-1, -1), 4),
        ("RIGHTPADDING", (0, 0), (-1, -1), 4),
    ])))
    story.append(PageBreak())


def build_journal():
    story = title_story("A spacious mathematical field notebook for sketching, modeling, revising, and explaining.")
    story += section("Journal Norms", "Rough work belongs here.", "This journal is not a worksheet packet. It is the visible record of your group's evolving mathematical thinking.")
    story.append(card_grid([
        ("Value uncertainty", "A strong investigation often begins confused. Label assumptions, questions, disagreements, and places where a model stops making sense.", INK, WHITE),
        ("Revise visibly", "Cross out, annotate, draw arrows, compare versions, and explain why your thinking changed.", INK, WHITE),
        ("Models are arguments", "A model is not an answer. It is a claim built from assumptions, evidence, representations, and limits.", INK, WHITE),
    ], 3))
    story.append(spacer(12))
    story.append(card("Field note style", "Messy is acceptable. Hidden thinking is not. The strongest journals show graph, algebra, context, units, limitations, and revision together.", INK, WHITE))
    story.append(PageBreak())

    story += section("Use Model", "Shared field notebook. Individual accountability.", "The project is collaborative, but grading is individual. The journal and individual evidence serve different purposes.")
    story.append(card_grid([
        ("Physical group journal", GROUP_JOURNAL_MODEL, PURPLE, PURPLE_SOFT),
        ("Individual digital evidence", INDIVIDUAL_EVIDENCE_MODEL, BLUE, BLUE_SOFT),
        ("How they work together", "The group journal preserves the messy investigation trail. Individual evidence shows how each student reasoned, interpreted, revised, and contributed mathematically inside that shared work.", GREEN, GREEN_SOFT),
    ], 3))
    story.append(spacer(16))
    story.append(card("Field notebook rule", "Clean pages are not the goal. Crossed-out thinking belongs here. Strong investigations often become more nuanced, not more certain.", INK, WHITE))
    story.append(PageBreak())

    story += section("Rubric Alignment", "Use the journal to make standards visible.", "The rubric does not only score final answers. It looks for reasoning, interpretation, revision, limitations, precision, and flexible thinking across the investigation.")
    story.append(card_grid([(f"{n}. {title}", kind, accent, fill) for n, title, kind, accent, fill, descs in RUBRIC_STANDARDS], 3))
    story.append(PageBreak())

    for class_num, title, question, accent, pages in PHASES:
        story += section(class_num, title, question)
        story.append(card("Field notebook expectation", "Use this section as a place for rough work, uncertainty, recalculation, and revision. Let the thinking stay visible.", accent, WHITE))
        story.append(PageBreak())
        for label, prompt, grid in pages:
            journal_page(story, label, prompt, "student thinking space", grid)
        if class_num in ["Class 2", "Class 3", "Class 4"]:
            story += section(f"{class_num} Model Comparison Studio", "Changing assumptions creates different mathematical realities.", "Choose model versions that are meaningful for your context. Which model is more useful, more realistic, or more honest about uncertainty? Which model hides constraints, exaggerates certainty, or breaks first?")
            story.append(Table([[Workspace("version A", 245, True), Workspace("version B", 245, True), Workspace("version C", 245, True)]], colWidths=[(PAGE_W - 100) / 3] * 3, style=TableStyle([
                ("VALIGN", (0, 0), (-1, -1), "TOP"),
                ("LEFTPADDING", (0, 0), (-1, -1), 5),
                ("RIGHTPADDING", (0, 0), (-1, -1), 5),
            ])))
            story.append(PageBreak())

    d = doc(OUT / "investigation-journal.pdf", "Enhanced Investigation Journal")
    d.build(story, onFirstPage=footer, onLaterPages=footer)


def docx_color(hex_color):
    raw = hex_color.replace("#", "")
    return RGBColor(int(raw[0:2], 16), int(raw[2:4], 16), int(raw[4:6], 16))


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill.replace("#", ""))
    tc_pr.append(shd)


def set_cell_border(cell, color="DBE2EA", size="8"):
    tc_pr = cell._tc.get_or_add_tcPr()
    borders = tc_pr.first_child_found_in("w:tcBorders")
    if borders is None:
        borders = OxmlElement("w:tcBorders")
        tc_pr.append(borders)
    for edge in ("top", "left", "bottom", "right"):
        tag = "w:{}".format(edge)
        element = borders.find(qn(tag))
        if element is None:
            element = OxmlElement(tag)
            borders.append(element)
        element.set(qn("w:val"), "single")
        element.set(qn("w:sz"), size)
        element.set(qn("w:space"), "0")
        element.set(qn("w:color"), color)


def set_cell_margins(cell, margin=160):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for side in ("top", "left", "bottom", "right"):
        node = tc_mar.find(qn(f"w:{side}"))
        if node is None:
            node = OxmlElement(f"w:{side}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(margin))
        node.set(qn("w:type"), "dxa")


def style_docx(document):
    sec = document.sections[0]
    sec.orientation = WD_ORIENT.LANDSCAPE
    sec.page_width = Inches(11)
    sec.page_height = Inches(8.5)
    sec.left_margin = Inches(0.6)
    sec.right_margin = Inches(0.6)
    sec.top_margin = Inches(0.5)
    sec.bottom_margin = Inches(0.5)

    styles_doc = document.styles
    normal = styles_doc["Normal"]
    normal.font.name = "Arial"
    normal.font.size = Pt(10.5)
    normal.font.color.rgb = docx_color("#334155")
    normal.paragraph_format.line_spacing = 1.15
    normal.paragraph_format.space_after = Pt(6)

    for name, size, color, before, after in [
        ("Title", 34, "#111827", 0, 14),
        ("Heading 1", 22, "#111827", 12, 8),
        ("Heading 2", 15, "#111827", 10, 6),
        ("Heading 3", 12, "#111827", 8, 4),
    ]:
        st = styles_doc[name]
        st.font.name = "Arial"
        st.font.bold = True
        st.font.size = Pt(size)
        st.font.color.rgb = docx_color(color)
        st.paragraph_format.space_before = Pt(before)
        st.paragraph_format.space_after = Pt(after)


def add_docx_para(document, text, style=None, bold=False, color=None):
    paragraph = document.add_paragraph(style=style)
    run = paragraph.add_run(text)
    run.bold = bold
    if color:
        run.font.color.rgb = docx_color(color)
    return paragraph


def add_docx_card(parent, title, body, fill="#FFFFFF", accent="#111827"):
    table = parent.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    cell = table.cell(0, 0)
    set_cell_shading(cell, fill)
    set_cell_border(cell)
    set_cell_margins(cell, 190)
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.TOP
    p1 = cell.paragraphs[0]
    r1 = p1.add_run(title)
    r1.bold = True
    r1.font.size = Pt(12)
    r1.font.color.rgb = docx_color(accent)
    p2 = cell.add_paragraph(body)
    p2.paragraph_format.space_after = Pt(0)
    return table


def add_docx_card_grid(document, items, cols=3):
    rows = (len(items) + cols - 1) // cols
    table = document.add_table(rows=rows, cols=cols)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    idx = 0
    for r in range(rows):
        for c in range(cols):
            cell = table.cell(r, c)
            set_cell_margins(cell, 160)
            if idx < len(items):
                title, body, fill, accent = items[idx]
                set_cell_shading(cell, fill)
                set_cell_border(cell)
                p1 = cell.paragraphs[0]
                p1.paragraph_format.space_after = Pt(3)
                run = p1.add_run(title)
                run.bold = True
                run.font.color.rgb = docx_color(accent)
                run.font.size = Pt(11.5)
                p2 = cell.add_paragraph(body)
                p2.paragraph_format.space_after = Pt(0)
            else:
                set_cell_border(cell, "FFFFFF", "0")
            idx += 1
    document.add_paragraph()
    return table


def add_docx_workspace(document, label, prompt, grid=False):
    add_docx_para(document, label, "Heading 2")
    add_docx_card(document, "Prompt", prompt, "#F8FAFC", "#111827")
    table = document.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    cell = table.cell(0, 0)
    set_cell_shading(cell, "#FFFFFF")
    set_cell_border(cell, "AEBBCC", "10")
    set_cell_margins(cell, 220)
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.TOP
    cell.paragraphs[0].add_run("Student thinking space").bold = True
    for _ in range(13 if grid else 11):
        p_blank = cell.add_paragraph("□ " if grid else " ")
        p_blank.paragraph_format.space_after = Pt(9)
    add_docx_card(document, "Reflection prompt", "What changed in your thinking? What assumption is doing the most work? What evidence would make this stronger?", "#F8FAFC", "#111827")


def build_journal_docx():
    document = Document()
    style_docx(document)
    add_docx_para(document, "ALGEBRA 2 CUMULATIVE SEMESTER FINAL PROJECT", bold=True, color="#64748B")
    title = document.add_paragraph(style="Title")
    title.add_run("Systems Under Pressure").bold = True
    add_docx_para(document, "The Mathematics of Food Systems", "Heading 1")
    add_docx_card(document, '"All models are wrong, but some are useful."', "George Box", "#FFFFFF", "#111827")
    add_docx_para(document, "A collaborative mathematical field notebook for sketching, modeling, revising, annotating, and explaining.", None)
    document.add_page_break()

    add_docx_para(document, "Journal Norms", "Heading 1")
    add_docx_card_grid(document, [
        ("Value uncertainty", "A strong investigation often begins confused. Label assumptions, questions, disagreements, and places where a model stops making sense.", "#FFFFFF", "#111827"),
        ("Revise visibly", "Cross out, annotate, draw arrows, compare versions, and explain why your thinking changed.", "#FFFFFF", "#111827"),
        ("Models are arguments", "A model is not an answer. It is a claim built from assumptions, evidence, representations, and limits.", "#FFFFFF", "#111827"),
    ], 3)
    add_docx_card(document, "Field note style", "Messy is acceptable. Hidden thinking is not. The strongest journals show graph, algebra, context, units, limitations, and revision together.", "#FFFFFF", "#111827")
    document.add_page_break()

    add_docx_para(document, "Shared field notebook. Individual accountability.", "Heading 1")
    add_docx_para(document, "The project is collaborative, but grading is individual. The journal and individual evidence serve different purposes.")
    add_docx_card_grid(document, [
        ("Physical group journal", GROUP_JOURNAL_MODEL, "#F6EFFF", "#7E22CE"),
        ("Individual digital evidence", INDIVIDUAL_EVIDENCE_MODEL, "#EDF4FF", "#2563EB"),
        ("How they work together", "The group journal preserves the messy investigation trail. Individual evidence shows how each student reasoned, interpreted, revised, and contributed mathematically inside that shared work.", "#EDF8F1", "#15803D"),
    ], 3)
    add_docx_card(document, "Field notebook rule", "Clean pages are not the goal. Crossed-out thinking belongs here. Strong investigations often become more nuanced, not more certain.", "#FFFFFF", "#111827")
    document.add_page_break()

    add_docx_para(document, "Rubric Alignment", "Heading 1")
    add_docx_para(document, "Use the journal to make standards visible. The rubric looks for reasoning, interpretation, revision, limitations, precision, and flexible thinking across the investigation.")
    add_docx_card_grid(document, [(f"{n}. {title}", kind, "#F8FAFC", "#111827") for n, title, kind, accent, fill, descs in RUBRIC_STANDARDS], 3)
    document.add_page_break()

    for class_num, title_text, question, accent, pages in PHASES:
        add_docx_para(document, f"{class_num} — {title_text}", "Heading 1")
        add_docx_para(document, question, None)
        add_docx_card(document, "Field notebook expectation", "Use this section as a place for rough work, uncertainty, recalculation, and revision. Let the thinking stay visible.", "#FFFFFF", "#111827")
        document.add_page_break()
        for label, prompt, grid in pages:
            add_docx_workspace(document, label, prompt, grid)
            document.add_page_break()
        if class_num in ["Class 2", "Class 3", "Class 4"]:
            add_docx_para(document, f"{class_num} Model Comparison Studio", "Heading 1")
            add_docx_para(document, "Changing assumptions creates different mathematical realities.")
            add_docx_card(document, "Comparison purpose", "Choose model versions that are meaningful for your context. Which model is more useful, more realistic, or more honest about uncertainty? Which model hides constraints, exaggerates certainty, or breaks first?", "#F8FAFC", "#111827")
            table = document.add_table(rows=1, cols=3)
            for i, cell in enumerate(table.rows[0].cells):
                set_cell_border(cell, "AEBBCC", "10")
                set_cell_margins(cell, 180)
                cell.paragraphs[0].add_run(f"Version {chr(65 + i)}").bold = True
                for _ in range(12):
                    cell.add_paragraph(" ")
            document.add_page_break()

    path = OUT / "investigation-journal.docx"
    document.save(path)
    return path


def build_rubric():
    # The official rubric is maintained directly at downloads/rubrics.pdf.
    # Do not regenerate it; the uploaded PDF is the single source of truth.
    return None


if __name__ == "__main__":
    build_journal()
    build_journal_docx()
    print("Generated journal PDF and DOCX in downloads/. Rubric PDF is preserved as the official uploaded file.")
